"""发布管理（SCN-005 P01/P02/P03）测试义务。

设计事实源：docs/30 …/SCN-005 分支结果矩阵 + 04A UINV-16/17/18。
覆盖：候选门禁 / 索引准入与必填阻塞 / 模板 schema 拦截 / Markdown 生成与绑定 /
编辑影响分类 / 定稿裁定与条目修订回流 / docx 转换与格式 / 检查打回 / 发布基线 / 幂等。
"""
import json
import uuid
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Pt

import app.db.models  # noqa: F401  register tables
from app.api.schemas import (
    ConfirmBaselineCommand,
    DocIndexEntryRead,
    ExportCheckCommand,
    FinalizeMarkdownCommand,
    GenerateMarkdownCommand,
    ItemConfirmCommand,
    ManualFallbackCommand,
    MarkdownEditCommand,
    ReopenIndexCommand,
    SaveIndexCommand,
    StartDocxExportCommand,
)
from app.db.base import Base, make_engine, make_session_factory
from app.db.models import Material, Project, RequirementItem
from app.domain.errors import InvalidInput, NotFound, RejectedTransition
from app.repositories.publication import SqlPublicationRepository
from app.scripts.import_packaged_templates import import_packaged_templates
from app.services.publication import (
    DocumentOrchestrationService,
    ExportExecutionService,
)

RAW_TEXT = (
    "系统应支持将确认态需求条目按模板导出为 docx。"
    "导出耗时不超过五秒。"
    "系统必须部署在企业内网。"
    "系统应提供 OpenAPI 兼容接口。"
    "历史订单数据至少保留三年。"
)


@pytest.fixture()
def session():
    engine = make_engine("sqlite://")
    Base.metadata.create_all(engine)
    factory = make_session_factory(engine)
    s = factory()
    import_packaged_templates(s)
    s.commit()
    try:
        yield s
    finally:
        s.close()
        engine.dispose()


def _seed(session):
    """项目 + 支撑材料 + 确认态条目（功能/质量/约束/接口）+ 1 条待确认功能条目。"""
    p = Project(name="发布演示项目")
    session.add(p)
    session.flush()
    mat = Material(project_id=p.id, raw_text=RAW_TEXT, source_note="评审纪要 2026-06")
    session.add(mat)
    session.flush()

    def item(req_no, expression, req_type, status):
        r = RequirementItem(
            project_id=p.id, parse_result_ref=uuid.uuid4(), formation_context_ref=uuid.uuid4(),
            req_no=req_no, expression=expression, req_type=req_type, status=status,
            source_element_refs="[]",
        )
        session.add(r)
        session.flush()
        return str(r.id)

    seeded = {
        "project": str(p.id), "material": str(mat.id),
        "fr1": item("FR-001", "系统应支持将确认态需求条目按模板导出为 docx", "functional", "confirmed"),
        "fr2": item("FR-002", "系统应提供 OpenAPI 兼容接口", "interface", "confirmed"),
        "q1": item("NFR-001", "导出耗时不超过五秒", "quality", "confirmed"),
        "c1": item("CON-001", "系统必须部署在企业内网", "constraint", "confirmed"),
        "pending": item("FR-009", "历史订单数据至少保留三年", "functional", "pending_confirmation"),
    }
    session.commit()
    return seeded


def _svc(session) -> DocumentOrchestrationService:
    return DocumentOrchestrationService(SqlPublicationRepository(session))


def _export_svc(session) -> ExportExecutionService:
    return ExportExecutionService(SqlPublicationRepository(session))  # inline 转换


def _entries(w, with_quality=True):
    entries = [
        DocIndexEntryRead(section_key="requirements.functional", asset_type="requirement_item",
                          asset_ref=w["fr1"], order_no=0),
        DocIndexEntryRead(section_key="requirements.interface", asset_type="requirement_item",
                          asset_ref=w["fr2"], order_no=0),
        DocIndexEntryRead(section_key="overview.constraints", asset_type="requirement_item",
                          asset_ref=w["c1"], order_no=0),
        DocIndexEntryRead(section_key="appendix.materials", asset_type="material",
                          asset_ref=w["material"], order_no=0),
    ]
    if with_quality:
        entries.append(DocIndexEntryRead(
            section_key="requirements.quality", asset_type="requirement_item",
            asset_ref=w["q1"], order_no=0))
    return entries


def _save_index(session, w, **kwargs):
    svc = _svc(session)
    result = svc.save_content_index(SaveIndexCommand(
        project_ref=w["project"], coverage_scope="release-v0.1 发布范围",
        entries=kwargs.pop("entries", _entries(w)),
        operator_ref="U1", idempotency_key=f"idx-{uuid.uuid4()}", **kwargs,
    ))
    session.commit()
    return svc, result


def _ready_draft(session, w):
    _svc_, save = _save_index(session, w)
    assert save.status == "index_ready"
    svc = _svc(session)
    draft = svc.generate_markdown(GenerateMarkdownCommand(
        project_ref=w["project"], operator_ref="U1", idempotency_key=f"gen-{uuid.uuid4()}",
    ))
    session.commit()
    return svc, draft


def _finalized_draft(session, w):
    svc, draft = _ready_draft(session, w)
    result = svc.finalize_markdown(FinalizeMarkdownCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key=f"fin-{uuid.uuid4()}",
    ))
    assert result.status == "finalized"
    session.commit()
    return svc, draft


# ============================================================================
# E0/E1：候选门禁 + 索引编排（P01）
# ============================================================================

def test_workspace_candidates_only_confirmed_items(session):
    w = _seed(session)
    ws = _svc(session).read_workspace(w["project"])
    refs = {i.item_ref for i in ws.candidates.items}
    assert w["fr1"] in refs and w["q1"] in refs
    assert w["pending"] not in refs  # 候选池只列确认态（US-E1-04）
    assert ws.candidates.pending_item_count == 1
    assert ws.template.error is None
    assert any(s.key == "requirements.functional" and s.required for s in ws.template.sections)


def test_item_confirm_min_gate(session):
    w = _seed(session)
    svc = _svc(session)
    result = svc.confirm_item(ItemConfirmCommand(
        project_ref=w["project"], item_ref=w["pending"], operator_ref="U1",
        idempotency_key="c-1",
    ))
    assert result.item_status == "confirmed"
    replay = svc.confirm_item(ItemConfirmCommand(
        project_ref=w["project"], item_ref=w["pending"], operator_ref="U1",
        idempotency_key="c-2",
    ))
    assert replay.status == "confirmed"  # 状态幂等重放
    ws = svc.read_workspace(w["project"])
    assert w["pending"] in {i.item_ref for i in ws.candidates.items}


def test_save_index_ready_opens_markdown(session):
    w = _seed(session)
    _, result = _save_index(session, w)
    assert result.status == "index_ready"
    assert result.index_version == 1
    assert result.missing_list == []
    ws = _svc(session).read_workspace(w["project"])
    assert ws.document.status == "index_ready"
    assert all(s.satisfied for s in ws.slot_status if s.required)


def test_save_index_missing_required_blocks(session):
    w = _seed(session)
    entries = [e for e in _entries(w) if e.section_key != "requirements.quality"]
    _, result = _save_index(session, w, entries=entries)
    assert result.status == "index_blocked"
    assert any(m.section_key == "requirements.quality" for m in result.missing_list)
    assert result.missing_list[0].rebuild_entry  # 补建入口
    with pytest.raises(RejectedTransition):  # 受阻索引不得生成 Markdown（US-E2-07）
        _svc(session).generate_markdown(GenerateMarkdownCommand(
            project_ref=w["project"], operator_ref="U1", idempotency_key="g-1",
        ))


def test_save_index_rejects_unconfirmed_item(session):
    w = _seed(session)
    entries = _entries(w) + [DocIndexEntryRead(
        section_key="requirements.functional", asset_type="requirement_item",
        asset_ref=w["pending"], order_no=9,
    )]
    with pytest.raises(RejectedTransition):  # 门禁不因模板必填而降低（US-E1-04）
        _save_index(session, w, entries=entries)


def test_save_index_rejects_slot_type_mismatch(session):
    w = _seed(session)
    entries = [DocIndexEntryRead(
        section_key="requirements.quality", asset_type="requirement_item",
        asset_ref=w["fr1"], order_no=0,  # 功能条目投质量槽位
    )]
    with pytest.raises(InvalidInput):
        _save_index(session, w, entries=entries)


def test_unregistered_template_blocks_index(session):
    w = _seed(session)
    _, result = _save_index(session, w, template_ref="srs-broken-v1", entries=[])
    assert result.status == "index_blocked"
    assert "template_registry" in (result.blocked_reason or "")
    ws = _svc(session).read_workspace(w["project"], template_ref="srs-broken-v1")
    assert ws.template.error is not None  # 不用默认结构或本地文件冒充（US-E1-05）


def test_optional_slot_empty_does_not_block(session):
    w = _seed(session)
    entries = [e for e in _entries(w) if e.section_key not in ("appendix.materials", "requirements.data")]
    _, result = _save_index(session, w, entries=entries)
    assert result.status == "index_ready"  # 可选缺失不阻塞（US-E1-06）


# ============================================================================
# E2：Markdown 生成、编辑影响、定稿（P02）
# ============================================================================

def test_generate_markdown_structure_and_bindings(session):
    w = _seed(session)
    _, draft = _ready_draft(session, w)
    assert "# 1 引言" in draft.content
    assert "## 3.1 功能需求" in draft.content
    assert "**FR-001**" in draft.content
    kinds = {b.kind for b in draft.source_bindings}
    assert {"heading", "boilerplate", "item", "material"} <= kinds
    ws = _svc(session).read_workspace(w["project"])
    assert ws.document.status == "markdown_draft"


def test_regenerate_creates_new_version(session):
    w = _seed(session)
    svc, draft1 = _ready_draft(session, w)
    draft2 = svc.generate_markdown(GenerateMarkdownCommand(
        project_ref=w["project"], operator_ref="U1", idempotency_key="g-2",
    ))
    session.commit()
    assert draft2.version_no == draft1.version_no + 1
    old = SqlPublicationRepository(session).get_draft(draft1.draft_ref)
    assert old.status == "superseded"


def test_doc_expression_edit_then_finalize(session):
    w = _seed(session)
    svc, draft = _ready_draft(session, w)
    edited = draft.content.replace("预期读者包括", "本说明的预期读者包括")
    result = svc.record_edit(MarkdownEditCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref, content=edited, operator_ref="U1",
    ))
    session.commit()
    assert [p.impact for p in result.patches] == ["doc_expression"]
    assert result.can_finalize
    fin = svc.finalize_markdown(FinalizeMarkdownCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="f-1",
    ))
    session.commit()
    assert fin.status == "finalized"
    row = SqlPublicationRepository(session).get_draft(draft.draft_ref)
    assert row.status == "finalized" and row.can_export
    assert "本说明的预期读者包括" in row.content  # 纯表达微调进入定稿
    ws = svc.read_workspace(w["project"])
    assert ws.document.status == "markdown_finalized"


def test_confirmed_item_edit_requires_reflow(session):
    w = _seed(session)
    svc, draft = _ready_draft(session, w)
    edited = draft.content.replace(
        "系统应支持将确认态需求条目按模板导出为 docx",
        "系统应支持将确认态需求条目按模板一键导出为 docx 与 PDF",
    )
    result = svc.record_edit(MarkdownEditCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref, content=edited, operator_ref="U1",
    ))
    session.commit()
    assert result.pending_item_refs == [w["fr1"]]

    # 未确认清单 → 不回流不定稿（US-E2-05 的"拒绝"即不带 confirm_reflow 重试/退回）
    fin1 = svc.finalize_markdown(FinalizeMarkdownCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="f-2",
    ))
    assert fin1.status == "pending_item_confirmation"
    assert len(fin1.pending_items) == 1
    old = SqlPublicationRepository(session).get_item(w["fr1"])
    assert old.status == "confirmed"  # 旧确认态不变

    # 确认清单 → 回流生成新的待确认条目；当前稿等待收束不可定稿（US-E2-04）
    fin2 = svc.finalize_markdown(FinalizeMarkdownCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref, confirm_reflow=True,
        operator_ref="U1", idempotency_key="f-3",
    ))
    session.commit()
    assert fin2.status == "item_revision_reflowed"
    assert len(fin2.reflowed_item_refs) == 1
    repo = SqlPublicationRepository(session)
    new_item = repo.get_item(fin2.reflowed_item_refs[0])
    assert new_item.status == "pending_confirmation"
    assert new_item.version_no == 2
    assert "PDF" in new_item.expression
    assert repo.get_item(w["fr1"]).status == "confirmed"  # 不原地覆盖
    assert repo.get_item(w["fr1"]).expression.endswith("docx")
    row = repo.get_draft(draft.draft_ref)
    assert row.status == "awaiting_item_revision" and not row.can_export


def test_no_source_fact_blocks_finalize(session):
    w = _seed(session)
    svc, draft = _ready_draft(session, w)
    edited = draft.content.replace(
        "# 附录A 支撑材料",
        "系统还应支持区块链存证与量子加密传输能力。\n\n# 附录A 支撑材料",
    )
    result = svc.record_edit(MarkdownEditCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref, content=edited, operator_ref="U1",
    ))
    session.commit()
    assert "no_source_fact" in [p.impact for p in result.patches]
    assert not result.can_finalize
    fin = svc.finalize_markdown(FinalizeMarkdownCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="f-4",
    ))
    assert fin.status == "blocked"
    assert any("无法支撑" in r for r in fin.block_reasons)


def test_heading_edit_classified_index_structure(session):
    w = _seed(session)
    svc, draft = _ready_draft(session, w)
    edited = draft.content.replace("## 3.1 功能需求", "## 3.1 核心功能")
    result = svc.record_edit(MarkdownEditCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref, content=edited, operator_ref="U1",
    ))
    assert "index_structure" in [p.impact for p in result.patches]
    assert not result.can_finalize  # 章节结构调整必须回 P01


def test_reopen_index_supersedes_draft(session):
    w = _seed(session)
    svc, draft = _ready_draft(session, w)
    doc = svc.reopen_index(ReopenIndexCommand(project_ref=w["project"], operator_ref="U1"))
    session.commit()
    assert doc.status == "index_ready"
    row = SqlPublicationRepository(session).get_draft(draft.draft_ref)
    assert row.status == "superseded"  # 原稿标记需重新生成（US-E1-07）


# ============================================================================
# E3：docx 导出、检查、发布基线（P03）
# ============================================================================

def test_export_rejected_when_not_finalized(session):
    w = _seed(session)
    _, draft = _ready_draft(session, w)
    result = _export_svc(session).start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-1",
    ))
    assert result.status == "rejected_precheck"  # 未定稿不可导出（US-E3-03）
    assert "P02" in result.next_action


def test_export_succeeds_and_docx_format_correct(session, tmp_path):
    from app import config
    object.__setattr__(config.settings, "export_dir", str(tmp_path))  # frozen dataclass 测试注入
    w = _seed(session)
    _, draft = _finalized_draft(session, w)
    exp = _export_svc(session)
    result = exp.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-2",
    ))
    session.commit()
    assert result.status == "submitted"
    repo = SqlPublicationRepository(session)
    export = repo.get_export(result.export_ref)
    assert export.status == "succeeded"
    assert export.file_path and export.file_path.endswith(".docx")

    # docx 格式断言（US-E3-02）：标题层级 + 正文段首行缩进 2 字符 + 字体绑定
    doc = DocxDocument(export.file_path)
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any(p.text == "1 引言" and p.style.name == "Heading 1" for p in headings)
    assert any(p.text == "3.1 功能需求" and p.style.name == "Heading 2" for p in headings)
    bodies = [p for p in doc.paragraphs
              if p.style.name == "Normal" and p.text.strip() and "FR-001" not in p.text
              and p.paragraph_format.first_line_indent is not None]
    assert bodies, "正文段落必须存在且带首行缩进"
    for p in bodies:
        assert p.paragraph_format.first_line_indent == Pt(24)  # 2 字符 × 小四 12pt
    fr = next(p for p in doc.paragraphs if "FR-001" in p.text)
    assert fr.paragraph_format.first_line_indent == Pt(24)
    assert any(r.font.bold for r in fr.runs)  # 条目编号加粗
    # 幂等重放
    replay = exp.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-2",
    ))
    assert replay.export_ref == result.export_ref


def test_export_conversion_failure_docks(session, tmp_path):
    from app import config
    object.__setattr__(config.settings, "export_dir", str(tmp_path))  # frozen dataclass 测试注入
    w = _seed(session)
    svc, draft = _ready_draft(session, w)
    edited = draft.content + "\n<!--convert-fail-->\n"
    svc.record_edit(MarkdownEditCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref, content=edited, operator_ref="U1",
    ))
    fin = svc.finalize_markdown(FinalizeMarkdownCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="f-5",
    ))
    assert fin.status == "finalized"
    session.commit()
    result = _export_svc(session).start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-3",
    ))
    session.commit()
    repo = SqlPublicationRepository(session)
    export = repo.get_export(result.export_ref)
    assert export.status == "failed"
    assert export.failure_reason  # 失败原因留存；Markdown/索引不变（US-E3-04）
    assert repo.get_draft(draft.draft_ref).status == "finalized"

    # 转换失败后可登记人工降级（US-E3-08）
    fb = _export_svc(session).register_manual_fallback(ManualFallbackCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        reason="转换能力不可用，线下人工排版", operator_ref="U1", idempotency_key="mf-1",
    ))
    session.commit()
    assert fb.manual_fallback and fb.status == "manual_fallback"
    baseline = _export_svc(session).confirm_baseline(ConfirmBaselineCommand(
        project_ref=w["project"], export_ref=fb.export_ref,
        operator_ref="U1", idempotency_key="b-0",
    ))
    session.commit()
    b = repo.get_baseline(baseline.baseline_ref)
    assert b.manual_fallback  # 基线带人工降级标记


def test_manual_fallback_rejected_without_failure(session):
    w = _seed(session)
    _, draft = _finalized_draft(session, w)
    with pytest.raises(RejectedTransition):
        _export_svc(session).register_manual_fallback(ManualFallbackCommand(
            project_ref=w["project"], draft_ref=draft.draft_ref,
            reason="想跳过转换", operator_ref="U1", idempotency_key="mf-2",
        ))


def test_check_rejected_blocks_baseline(session, tmp_path):
    from app import config
    object.__setattr__(config.settings, "export_dir", str(tmp_path))  # frozen dataclass 测试注入
    w = _seed(session)
    _, draft = _finalized_draft(session, w)
    exp = _export_svc(session)
    result = exp.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-4",
    ))
    session.commit()
    checked = exp.report_check(ExportCheckCommand(
        project_ref=w["project"], export_ref=result.export_ref, passed=False,
        note="目录编号错位", operator_ref="U1",
    ))
    session.commit()
    assert checked.status == "check_rejected"  # US-E3-05
    baseline = exp.confirm_baseline(ConfirmBaselineCommand(
        project_ref=w["project"], export_ref=result.export_ref,
        operator_ref="U1", idempotency_key="b-1",
    ))
    assert baseline.status == "rejected_precheck"


def test_confirm_baseline_freezes_snapshot(session, tmp_path):
    from app import config
    object.__setattr__(config.settings, "export_dir", str(tmp_path))  # frozen dataclass 测试注入
    w = _seed(session)
    _, draft = _finalized_draft(session, w)
    exp = _export_svc(session)
    result = exp.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-5",
    ))
    session.commit()
    repo = SqlPublicationRepository(session)

    # 导出成功≠发布：未确认前无基线（US-E3-07）
    doc = repo.get_document(w["project"])
    assert repo.baseline_of(str(doc.id)) is None

    baseline = exp.confirm_baseline(ConfirmBaselineCommand(
        project_ref=w["project"], export_ref=result.export_ref,
        note="验收通过", operator_ref="U-审定人", idempotency_key="b-2",
    ))
    session.commit()
    assert baseline.status == "confirmed"
    b = repo.get_baseline(baseline.baseline_ref)
    assert b.confirmed_by == "U-审定人"
    assert b.index_version == 1
    assert str(b.draft_ref) == draft.draft_ref
    refs = json.loads(b.asset_refs)
    assert any(w["fr1"] in r for r in refs)  # 冻结源资产版本引用（US-E3-06）
    assert repo.get_export(result.export_ref).status == "baseline_confirmed"
    ws = _svc(session).read_workspace(w["project"])
    assert ws.document.status == "baseline_published"
    assert ws.baseline is not None

    # 基线确认幂等重放
    replay = exp.confirm_baseline(ConfirmBaselineCommand(
        project_ref=w["project"], export_ref=result.export_ref,
        operator_ref="U-审定人", idempotency_key="b-3",
    ))
    assert replay.status == "confirmed" and replay.baseline_ref == baseline.baseline_ref


# ============================================================================
# 跨切面不变式：派生制品不反向覆盖正式资产
# ============================================================================

def test_publication_flow_never_mutates_source_assets(session, tmp_path):
    from app import config
    object.__setattr__(config.settings, "export_dir", str(tmp_path))  # frozen dataclass 测试注入
    w = _seed(session)
    repo = SqlPublicationRepository(session)
    before = {r: (repo.get_item(r).expression, repo.get_item(r).status)
              for r in (w["fr1"], w["fr2"], w["q1"], w["c1"])}
    mat_before = repo.get_materials([w["material"]])[0].raw_text

    _, draft = _finalized_draft(session, w)
    exp = _export_svc(session)
    result = exp.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-6",
    ))
    exp.confirm_baseline(ConfirmBaselineCommand(
        project_ref=w["project"], export_ref=result.export_ref,
        operator_ref="U1", idempotency_key="b-4",
    ))
    session.commit()

    for r, (expr, status) in before.items():
        item = repo.get_item(r)
        assert (item.expression, item.status) == (expr, status)
    assert repo.get_materials([w["material"]])[0].raw_text == mat_before


def test_unknown_refs_raise_not_found(session):
    w = _seed(session)
    svc = _svc(session)
    with pytest.raises(NotFound):
        svc.confirm_item(ItemConfirmCommand(
            project_ref=w["project"], item_ref=str(uuid.uuid4()),
            operator_ref="U1", idempotency_key="x-1",
        ))
    with pytest.raises(NotFound):
        _save_index(session, w, entries=[DocIndexEntryRead(
            section_key="requirements.functional", asset_type="requirement_item",
            asset_ref=str(uuid.uuid4()), order_no=0,
        )])


def test_docx_lands_in_configured_export_dir(tmp_path):
    """设置页保存的导出目录真正决定 docx 落盘位置（T20260724 走查发现：此前只读 env）。

    env 指向 env-exports、配置库存 saved-exports，导出后文件必须出现在后者。
    """
    from app import config
    from app.api.schemas import ConfigSaveCommand
    from app.repositories.agent_run import SqlAgentRunRepository
    from app.services.config_registry import ConfigRegistryService
    from app.services.publication import run_docx_export_judgement

    env_dir, saved_dir = tmp_path / "env-exports", tmp_path / "saved-exports"
    object.__setattr__(config.settings, "export_dir", str(env_dir))
    engine = make_engine(f"sqlite:///{tmp_path}/pub-dir.db")
    Base.metadata.create_all(engine)
    factory = make_session_factory(engine)
    session = factory()
    import_packaged_templates(session)
    ConfigRegistryService(session).save_domain(
        "export",
        ConfigSaveCommand(values={"export_dir": str(saved_dir)}, secrets={}, operator_ref="U1"),
    )
    session.commit()
    w = _seed(session)
    _, draft = _finalized_draft(session, w)

    svc = ExportExecutionService(
        SqlPublicationRepository(session),
        agent_runs=SqlAgentRunRepository(session),
        enqueue=lambda export_ref, run_id: run_docx_export_judgement(
            SqlPublicationRepository(session), export_ref
        ),
    )
    result = svc.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-dir1",
    ))
    session.commit()

    export = SqlPublicationRepository(session).get_export(result.export_ref)
    assert export.status == "succeeded"
    assert Path(export.file_path).parent == saved_dir
    assert Path(export.file_path).exists()
    assert not env_dir.exists()  # env 目录一个文件都不该产生
    session.close()
    engine.dispose()


def test_docx_lands_in_env_dir_when_nothing_saved(tmp_path):
    """没在设置页存过导出目录（全新部署的常态）：落点必须是 env 指定的那个目录。

    上面那条姊妹用例证明保存值生效，这条钉住回落分支——生产调用形态（不传 base）走的正是它，
    而 convert_markdown_to_docx 会自动建目录，落点写到哪里都不会报错，不断言父目录就发现不了。
    """
    from app import config
    from app.repositories.agent_run import SqlAgentRunRepository
    from app.services.publication import run_docx_export_judgement

    env_dir = tmp_path / "env-exports"
    object.__setattr__(config.settings, "export_dir", str(env_dir))
    engine = make_engine(f"sqlite:///{tmp_path}/pub-envdir.db")
    Base.metadata.create_all(engine)
    factory = make_session_factory(engine)
    session = factory()
    import_packaged_templates(session)
    session.commit()
    w = _seed(session)
    _, draft = _finalized_draft(session, w)

    svc = ExportExecutionService(
        SqlPublicationRepository(session),
        agent_runs=SqlAgentRunRepository(session),
        enqueue=lambda export_ref, run_id: run_docx_export_judgement(
            SqlPublicationRepository(session), export_ref
        ),
    )
    result = svc.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-dir2",
    ))
    session.commit()

    export = SqlPublicationRepository(session).get_export(result.export_ref)
    assert export.status == "succeeded"
    assert Path(export.file_path).parent == env_dir  # 无配置行 → env，不是进程当前目录
    session.close()
    engine.dispose()


def test_export_commits_before_enqueue(tmp_path):
    """落库后再入队：inline/worker 任务用独立 session，必须能读到导出记录。"""
    from app import config
    from app.repositories.agent_run import SqlAgentRunRepository
    from app.services.publication import run_docx_export_judgement

    object.__setattr__(config.settings, "export_dir", str(tmp_path))
    engine = make_engine(f"sqlite:///{tmp_path}/pub.db")
    Base.metadata.create_all(engine)
    factory = make_session_factory(engine)
    session = factory()
    import_packaged_templates(session)
    session.commit()
    w = _seed(session)
    _, draft = _finalized_draft(session, w)

    seen = {}

    def fake_enqueue(export_ref: str, run_id: str) -> None:
        other = factory()  # 模拟 worker：独立 session
        try:
            run_docx_export_judgement(SqlPublicationRepository(other), export_ref)
            other.commit()
            seen["status"] = SqlPublicationRepository(other).get_export(export_ref).status
        finally:
            other.close()

    svc = ExportExecutionService(
        SqlPublicationRepository(session),
        agent_runs=SqlAgentRunRepository(session), enqueue=fake_enqueue,
    )
    result = svc.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-q1",
    ))
    session.commit()
    assert result.status == "submitted" and result.agent_run_ref
    assert seen["status"] == "succeeded"
    session.close()
    engine.dispose()


def test_export_unexpected_error_docks_not_stuck(session, tmp_path, monkeypatch):
    """未预期异常（非 ConversionError/TemplateError）也必须落 failed，绝不遗留 converting。

    回归缺陷：过窄的 except 让写盘/样式等异常外抛 → worker rollback → 导出行永远 converting、
    界面永远『转换中』。修复后任何异常都落终态且不写异常原文（硬规则 8）。
    """
    from app import config
    from app.services import publication as pub

    object.__setattr__(config.settings, "export_dir", str(tmp_path))
    w = _seed(session)
    _, draft = _finalized_draft(session, w)

    def boom(*args, **kwargs):
        raise OSError("disk full / permission denied at /secret/path")

    monkeypatch.setattr(pub, "convert_markdown_to_docx", boom)
    result = _export_svc(session).start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-unexpected",
    ))
    session.commit()
    export = SqlPublicationRepository(session).get_export(result.export_ref)
    assert export.status == "failed"  # 不再遗留 converting
    assert export.failure_reason  # 有脱敏通用原因
    assert "disk full" not in export.failure_reason  # 不落异常原文（硬规则 8）


def test_export_dedup_inflight_conversion(session, tmp_path):
    """在途去重：同一定稿已有 converting 导出时，换幂等键再次提交也复用在途、不重复入队。"""
    from app.repositories.agent_run import SqlAgentRunRepository

    w = _seed(session)
    _, draft = _finalized_draft(session, w)
    # noop enqueue：模拟已入队但 worker 未执行 → 导出保持 converting
    first_svc = ExportExecutionService(
        SqlPublicationRepository(session),
        agent_runs=SqlAgentRunRepository(session), enqueue=lambda e, r: None,
    )
    first = first_svc.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-inflight-1",
    ))
    session.commit()
    repo = SqlPublicationRepository(session)
    assert repo.get_export(first.export_ref).status == "converting"
    doc_ref = str(repo.get_export(first.export_ref).document_ref)

    enq_calls: list[str] = []
    second_svc = ExportExecutionService(
        SqlPublicationRepository(session),
        agent_runs=SqlAgentRunRepository(session),
        enqueue=lambda e, r: enq_calls.append(e),
    )
    second = second_svc.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-inflight-2",  # 不同幂等键
    ))
    session.commit()
    assert second.export_ref == first.export_ref  # 复用在途，非新建
    assert enq_calls == []  # 未重复入队
    assert len(repo.exports_of(doc_ref)) == 1  # 只有一条导出行


def test_read_workspace_reconciles_failed_agentrun(session, tmp_path):
    """读侧自愈：converting 导出其 AgentRun 已 failed → 读工作区时对账落 failed 并持久化。"""
    from app.repositories.agent_run import SqlAgentRunRepository

    w = _seed(session)
    _, draft = _finalized_draft(session, w)
    ar = SqlAgentRunRepository(session)
    svc = ExportExecutionService(
        SqlPublicationRepository(session), agent_runs=ar, enqueue=lambda e, r: None,
    )
    started = svc.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-recon-1",
    ))
    session.commit()
    run = ar.find_by_context(started.export_ref, "docx_export")
    assert run is not None
    ar.mark_failed(str(run.id), "boom")
    session.commit()

    ws = _svc(session).read_workspace(w["project"])
    exp = next(e for e in ws.exports if e.export_ref == started.export_ref)
    assert exp.status == "failed"  # 读侧对账已终态化
    assert SqlPublicationRepository(session).get_export(started.export_ref).status == "failed"


def test_read_workspace_reconciles_timed_out_export(session, tmp_path):
    """读侧自愈：converting 超时（worker 缺席场景）且无产物 → 读工作区时对账落 failed。"""
    from datetime import datetime, timedelta, timezone

    from app.repositories.agent_run import SqlAgentRunRepository

    w = _seed(session)
    _, draft = _finalized_draft(session, w)
    svc = ExportExecutionService(
        SqlPublicationRepository(session),
        agent_runs=SqlAgentRunRepository(session), enqueue=lambda e, r: None,
    )
    started = svc.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-recon-2",
    ))
    session.commit()
    repo = SqlPublicationRepository(session)
    row = repo.get_export(started.export_ref)
    row.created_at = datetime.now(timezone.utc) - timedelta(minutes=30)  # 拨到超时之前
    session.commit()

    ws = _svc(session).read_workspace(w["project"])
    exp = next(e for e in ws.exports if e.export_ref == started.export_ref)
    assert exp.status == "failed"


# ---- 29148 属性补齐：条目块属性行（确定性投影、空值不渲染）+ 文档约定静态段 ----

def test_item_block_renders_verification_and_priority_rows(session):
    w = _seed(session)
    row = session.get(RequirementItem, uuid.UUID(w["fr1"]))
    row.verification_method = "demonstration,analysis"
    row.verification_note = "导出的 docx 可在验收环境打开且样式合规"
    row.priority = "high"
    session.commit()
    _svc_, draft = _ready_draft(session, w)
    assert "| 验证方式与验收准则 | 演示、分析：导出的 docx 可在验收环境打开且样式合规 |" in draft.content
    assert "| 优先级 | 高 |" in draft.content
    # 空值不渲染：未补属性的条目块不出现这两行标签（q1/c1/fr2 均未设）
    assert draft.content.count("| 优先级 |") == 1
    assert draft.content.count("| 验证方式与验收准则 |") == 1


def test_markdown_contains_conventions_boilerplate(session):
    """模板静态"文档约定"段：验证方式四方法定义 + 优先级三级口径（零生成）。"""
    w = _seed(session)
    _svc_, draft = _ready_draft(session, w)
    assert "## 1.4 文档约定" in draft.content
    assert "测试=以可执行用例对照预期结果验证" in draft.content
    assert "高=本发布范围内不满足即不可交付" in draft.content
