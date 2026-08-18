"""发布管理·图表入索引 + 资产→文档片段追溯（04A §8 基线增补）测试义务。

覆盖：受控图表候选门禁 / 图表槽位准入与错槽拒绝 / Markdown 图表渲染与绑定 /
图表源码编辑阻断（OTHER_ASSET）/ 图表失效回退索引受阻 / docx 围栏等宽段 /
基线冻结图表版本引用 / asset-fragment 各业务态 / chart→document 派生承接边。
"""
import json
import uuid

import pytest
from docx import Document as DocxDocument

import app.db.models  # noqa: F401  register tables
from app.api.schemas import (
    ConfirmBaselineCommand,
    DocIndexEntryRead,
    FinalizeMarkdownCommand,
    GenerateMarkdownCommand,
    MarkdownEditCommand,
    SaveIndexCommand,
    StartDocxExportCommand,
)
from app.db.base import Base, make_engine, make_session_factory
from app.db.models import (
    Material,
    MaterialParseResult,
    Project,
    RequirementChart,
    RequirementElement,
    RequirementItem,
    TraceLink,
)
from app.domain.errors import InvalidInput, NotFound, RejectedTransition
from app.repositories.publication import SqlPublicationRepository
from app.scripts.import_packaged_templates import import_packaged_templates
from app.services.publication import DocumentOrchestrationService, ExportExecutionService

MERMAID_CODE = "flowchart TD\n  A[接收材料] --> B[知识抽取]\n  B --> C[条目形成]"


@pytest.fixture()
def session():
    engine = make_engine("sqlite://")
    Base.metadata.create_all(engine)
    factory = make_session_factory(engine)
    s = factory()
    import_packaged_templates(s)
    s.commit()
    try:
        yield s
    finally:
        s.close()
        engine.dispose()


def _seed(session):
    p = Project(name="图表发布演示")
    session.add(p)
    session.flush()
    mat = Material(project_id=p.id, raw_text="系统应支持导出 docx。导出耗时不超过五秒。",
                   source_note="评审纪要")
    session.add(mat)
    session.flush()
    parse_result = MaterialParseResult(
        project_id=p.id, material_ref=mat.id, context_ref=uuid.uuid4(), parse_status="parsed",
    )
    session.add(parse_result)
    session.flush()
    element = RequirementElement(
        project_id=p.id, parse_result_ref=parse_result.id,
        element_type="functional_requirement", content="系统应支持导出 docx",
        process_status="confirmed",
    )
    session.add(element)
    session.flush()

    def item(req_no, expression, req_type, element_refs=None, curation=None, boundary=None):
        r = RequirementItem(
            project_id=p.id, parse_result_ref=parse_result.id, formation_context_ref=uuid.uuid4(),
            req_no=req_no, expression=expression, req_type=req_type, status="confirmed",
            source_element_refs=json.dumps(element_refs or []),
            curation_note=curation, boundary_note=boundary,
        )
        session.add(r)
        session.flush()
        return str(r.id)

    def chart(title, status, fmt="mermaid", source=MERMAID_CODE, sources=None):
        c = RequirementChart(
            project_id=p.id, title=title, chart_kind="behavior", chart_type="flowchart",
            format=fmt, source_code=source, draft_version=2, status=status,
            source_refs=json.dumps(sources or []),
        )
        session.add(c)
        session.flush()
        return str(c.id)

    fr1 = item(
        "FR-001", "系统应支持将确认态需求条目按模板导出为 docx", "functional",
        element_refs=[str(element.id)],
        curation="合并两条同义表述，规范为 EARS 句式",
        boundary="仅覆盖 docx 导出；PDF 导出明确不在本条范围",
    )
    chart_ok = chart("形成管线流程图", "confirmed", sources=[fr1])
    session.add(TraceLink(
        project_id=p.id, relation_type="chart", status="effective",
        upstream_type="requirement_item", upstream_ref=uuid.UUID(fr1),
        downstream_type="chart", downstream_ref=uuid.UUID(chart_ok),
        initial_basis="测试种子：图表确认后正式确立",
    ))
    seeded = {
        "project": str(p.id), "material": str(mat.id), "fr1": fr1,
        "q1": item("NFR-001", "导出耗时不超过五秒", "quality"),
        "chart_ok": chart_ok,
        "chart_draft": chart("草稿态泳道图", "draft"),
        "chart_table": chart("状态迁移表", "confirmed", fmt="markdown_table",
                             source="| 状态 | 事件 |\n|---|---|\n| 草稿 | 提交 |"),
    }
    session.commit()
    return seeded


def _svc(session) -> DocumentOrchestrationService:
    return DocumentOrchestrationService(SqlPublicationRepository(session))


def _entries(w, chart_entries=None):
    entries = [
        DocIndexEntryRead(section_key="requirements.functional", asset_type="requirement_item",
                          asset_ref=w["fr1"], order_no=0),
        DocIndexEntryRead(section_key="requirements.quality", asset_type="requirement_item",
                          asset_ref=w["q1"], order_no=0),
    ]
    entries += chart_entries if chart_entries is not None else [
        DocIndexEntryRead(section_key="requirements.charts", asset_type="chart",
                          asset_ref=w["chart_ok"], order_no=0),
    ]
    return entries


def _save_index(session, w, entries):
    svc = _svc(session)
    result = svc.save_content_index(SaveIndexCommand(
        project_ref=w["project"], coverage_scope="release-v0.1",
        entries=entries, operator_ref="U1", idempotency_key=f"idx-{uuid.uuid4()}",
    ))
    session.commit()
    return svc, result


def _draft(session, w, entries=None):
    _, save = _save_index(session, w, entries if entries is not None else _entries(w))
    assert save.status == "index_ready"
    svc = _svc(session)
    draft = svc.generate_markdown(GenerateMarkdownCommand(
        project_ref=w["project"], operator_ref="U1", idempotency_key=f"gen-{uuid.uuid4()}",
    ))
    session.commit()
    return svc, draft


# ---- 候选池：受控图表门禁 ----

def test_candidates_include_only_confirmed_charts(session):
    w = _seed(session)
    ws = _svc(session).read_workspace(w["project"])
    refs = {c.chart_ref for c in ws.candidates.charts}
    assert w["chart_ok"] in refs and w["chart_table"] in refs
    assert w["chart_draft"] not in refs  # 仅受控（confirmed）入候选
    ok = next(c for c in ws.candidates.charts if c.chart_ref == w["chart_ok"])
    assert ok.source_count == 1 and ok.draft_version == 2 and ok.format == "mermaid"
    assert ws.candidates.trace_summary is not None
    assert any(s.key == "requirements.charts" and not s.required for s in ws.template.sections)


# ---- P01：图表入槽准入 ----

def test_chart_into_chart_slot_saves_ready(session):
    w = _seed(session)
    _, result = _save_index(session, w, _entries(w))
    assert result.status == "index_ready"
    ws = _svc(session).read_workspace(w["project"])
    entry = next(e for e in ws.index_entries if e.asset_type == "chart")
    assert entry.asset_version == "2"  # 冻结 draft_version


def test_chart_into_item_slot_rejected(session):
    w = _seed(session)
    entries = _entries(w, chart_entries=[DocIndexEntryRead(
        section_key="requirements.functional", asset_type="chart",
        asset_ref=w["chart_ok"], order_no=9,
    )])
    with pytest.raises(InvalidInput):
        _save_index(session, w, entries)


def test_item_into_chart_slot_rejected(session):
    w = _seed(session)
    entries = [DocIndexEntryRead(
        section_key="requirements.charts", asset_type="requirement_item",
        asset_ref=w["fr1"], order_no=0,
    )]
    with pytest.raises(InvalidInput):
        _save_index(session, w, entries)


def test_unconfirmed_chart_rejected(session):
    w = _seed(session)
    entries = _entries(w, chart_entries=[DocIndexEntryRead(
        section_key="requirements.charts", asset_type="chart",
        asset_ref=w["chart_draft"], order_no=0,
    )])
    with pytest.raises(RejectedTransition):
        _save_index(session, w, entries)


def test_unknown_chart_ref_not_found(session):
    w = _seed(session)
    entries = _entries(w, chart_entries=[DocIndexEntryRead(
        section_key="requirements.charts", asset_type="chart",
        asset_ref=str(uuid.uuid4()), order_no=0,
    )])
    with pytest.raises(NotFound):
        _save_index(session, w, entries)


# ---- P02：渲染、绑定与编辑阻断 ----

def test_markdown_renders_chart_fence_and_binding(session):
    w = _seed(session)
    _, draft = _draft(session, w)
    assert "## 3.5 需求图表" in draft.content
    assert "**图：形成管线流程图**（flowchart）" in draft.content
    assert "```mermaid" in draft.content
    assert MERMAID_CODE in draft.content
    chart_bindings = [b for b in draft.source_bindings if b.kind == "chart"]
    assert len(chart_bindings) == 1
    assert chart_bindings[0].asset_ref == w["chart_ok"]


def test_markdown_table_chart_rendered_unfenced(session):
    w = _seed(session)
    entries = _entries(w, chart_entries=[DocIndexEntryRead(
        section_key="requirements.charts", asset_type="chart",
        asset_ref=w["chart_table"], order_no=0,
    )])
    _, draft = _draft(session, w, entries)
    assert "| 状态 | 事件 |" in draft.content
    assert "```markdown_table" not in draft.content  # 表格图直接以 Markdown 表呈现


def test_chart_source_edit_blocks_finalize_as_other_asset(session):
    w = _seed(session)
    svc, draft = _draft(session, w)
    edited = draft.content.replace("B[知识抽取]", "B[知识抽取与自动确认]")
    result = svc.record_edit(MarkdownEditCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref, content=edited, operator_ref="U1",
    ))
    session.commit()
    assert "other_asset" in [p.impact for p in result.patches]
    assert not result.can_finalize
    fin = svc.finalize_markdown(FinalizeMarkdownCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="f-c1",
    ))
    assert fin.status == "blocked"
    assert any("本流程外正式资产" in r for r in fin.block_reasons)


def test_item_block_renders_statement_and_attribute_table(session):
    """条目块（29148 §5.2.8 属性化）：规范陈述 + 整理/边界说明 + 来源依据 + 关联图表。"""
    w = _seed(session)
    _, draft = _draft(session, w)
    assert "**FR-001**（功能需求 · v1 · 已确认）" in draft.content
    assert "\n系统应支持将确认态需求条目按模板导出为 docx\n" in draft.content
    assert "| 内容整理说明 | 合并两条同义表述，规范为 EARS 句式 |" in draft.content
    assert "| 条目边界说明 | 仅覆盖 docx 导出；PDF 导出明确不在本条范围 |" in draft.content
    assert "| 来源依据 | 评审纪要（材料 v1）（来源要素 1 项） |" in draft.content
    assert "| 关联图表 | 形成管线流程图 |" in draft.content
    # 无说明/来源/图表的条目不渲染空属性表
    q1_block_start = draft.content.index("**NFR-001**")
    q1_block = draft.content[q1_block_start:q1_block_start + 200]
    assert "| 属性 |" not in q1_block
    # 条目块仍是单一 item 绑定（编辑影响与片段追溯口径不变）
    fr1_bindings = [b for b in draft.source_bindings if b.kind == "item" and b.asset_ref == w["fr1"]]
    assert len(fr1_bindings) == 1


def test_attribute_row_edit_reflows_with_original_expression(session):
    """属性表行编辑 → 仍判 confirmed_item 回流，但不得把表格文本当新表达。"""
    w = _seed(session)
    svc, draft = _draft(session, w)
    edited = draft.content.replace(
        "| 条目边界说明 | 仅覆盖 docx 导出；PDF 导出明确不在本条范围 |",
        "| 条目边界说明 | 覆盖 docx 与 PDF 导出 |",
    )
    result = svc.record_edit(MarkdownEditCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref, content=edited, operator_ref="U1",
    ))
    session.commit()
    assert result.pending_item_refs == [w["fr1"]]
    fin = svc.finalize_markdown(FinalizeMarkdownCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref, confirm_reflow=True,
        operator_ref="U1", idempotency_key="f-attr",
    ))
    session.commit()
    assert fin.status == "item_revision_reflowed"
    repo = SqlPublicationRepository(session)
    new_item = repo.get_item(fin.reflowed_item_refs[0])
    # 表格行提取不出正文 → 回退原表达（不把 "| 条目边界说明 |…" 写成需求）
    assert new_item.expression == "系统应支持将确认态需求条目按模板导出为 docx"


def test_voided_chart_blocks_regeneration(session):
    w = _seed(session)
    svc, _ = _draft(session, w)
    chart = session.get(RequirementChart, uuid.UUID(w["chart_ok"]))
    chart.status = "voided"
    session.commit()
    with pytest.raises(RejectedTransition):
        svc.generate_markdown(GenerateMarkdownCommand(
            project_ref=w["project"], operator_ref="U1", idempotency_key="g-void",
        ))
    doc = SqlPublicationRepository(session).get_document(w["project"])
    assert doc.status == "index_blocked"


# ---- P03：docx 围栏 + 基线冻结 ----

def test_docx_renders_mermaid_fence_as_image(session, tmp_path):
    from app import config
    object.__setattr__(config.settings, "export_dir", str(tmp_path))
    w = _seed(session)
    svc, draft = _draft(session, w)
    fin = svc.finalize_markdown(FinalizeMarkdownCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="f-c2",
    ))
    assert fin.status == "finalized"
    session.commit()
    exp = ExportExecutionService(SqlPublicationRepository(session))
    result = exp.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-c1",
    ))
    session.commit()
    repo = SqlPublicationRepository(session)
    export = repo.get_export(result.export_ref)
    assert export.status == "succeeded"
    doc = DocxDocument(export.file_path)
    texts = [p.text for p in doc.paragraphs]
    assert not any("```" in t for t in texts)  # 围栏行不输出
    # 图形围栏本地栅格化为内嵌图片：源码不再作为文本外泄
    assert not any("A[接收材料] --> B[知识抽取]" in t for t in texts)
    assert doc.inline_shapes, "mermaid 图形应渲染为内嵌图片"
    assert any("PICTURE" in str(s.type) for s in doc.inline_shapes)
    # 条目属性表渲染为真 docx 表格（表头加粗、分隔行不输出、管道字符不外泄）
    assert doc.tables, "条目属性表应渲染为 docx 表格"
    attr_table = next(t for t in doc.tables if t.rows[0].cells[0].text == "属性")
    assert attr_table.rows[0].cells[1].text == "说明"
    body_cells = {c.text for row in attr_table.rows[1:] for c in row.cells}
    assert "内容整理说明" in body_cells and "关联图表" in body_cells
    assert not any("|" in t for t in texts)  # 表格行不落入正文段落

    # 基线冻结图表版本引用
    baseline = exp.confirm_baseline(ConfirmBaselineCommand(
        project_ref=w["project"], export_ref=result.export_ref,
        operator_ref="U1", idempotency_key="b-c1",
    ))
    session.commit()
    refs = json.loads(repo.get_baseline(baseline.baseline_ref).asset_refs)
    assert f"chart:{w['chart_ok']}@v2" in refs


# ---- asset-fragment：资产 → 文档片段追溯 ----

def test_fragment_before_orchestration_is_soft_empty(session):
    w = _seed(session)
    read = _svc(session).read_asset_fragment(w["project"], "requirement_item", w["fr1"])
    assert read.fragments == [] and read.document_ref is None
    assert "尚未进行文档编排" in read.next_action


def test_fragment_rejects_unsupported_asset_type(session):
    w = _seed(session)
    with pytest.raises(InvalidInput):
        _svc(session).read_asset_fragment(w["project"], "material", w["material"])


def test_fragment_after_index_before_markdown(session):
    w = _seed(session)
    _save_index(session, w, _entries(w))
    read = _svc(session).read_asset_fragment(w["project"], "chart", w["chart_ok"])
    assert read.in_current_index and read.draft_ref is None
    assert "尚未生成 Markdown" in read.next_action


def test_fragment_returns_exact_slices(session):
    w = _seed(session)
    _, draft = _draft(session, w)
    read = _svc(session).read_asset_fragment(w["project"], "chart", w["chart_ok"])
    assert read.in_current_index and read.draft_status == "draft"
    assert len(read.fragments) == 1
    frag = read.fragments[0]
    assert frag.section_key == "requirements.charts"
    assert frag.section_number == "3.5" and frag.section_title == "需求图表"
    assert frag.markdown.startswith("**图：形成管线流程图**")
    assert MERMAID_CODE in frag.markdown
    # 行区间与生成稿逐行一致
    lines = draft.content.splitlines()
    assert "\n".join(lines[frag.start_line:frag.end_line + 1]) == frag.markdown

    item_read = _svc(session).read_asset_fragment(w["project"], "requirement_item", w["fr1"])
    assert len(item_read.fragments) == 1
    assert "FR-001" in item_read.fragments[0].markdown


def test_fragment_not_arranged_asset_soft_empty(session):
    w = _seed(session)
    entries = _entries(w, chart_entries=[])  # 图表不入索引
    _draft(session, w, entries)
    read = _svc(session).read_asset_fragment(w["project"], "chart", w["chart_ok"])
    assert read.fragments == [] and not read.in_current_index
    assert "回索引编排页" in read.next_action


def test_fragment_marks_finalized_and_baseline(session, tmp_path):
    from app import config
    object.__setattr__(config.settings, "export_dir", str(tmp_path))
    w = _seed(session)
    svc, draft = _draft(session, w)
    svc.finalize_markdown(FinalizeMarkdownCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="f-c3",
    ))
    session.commit()
    read = svc.read_asset_fragment(w["project"], "chart", w["chart_ok"])
    assert read.draft_status == "finalized" and read.baseline_ref is None

    exp = ExportExecutionService(SqlPublicationRepository(session))
    result = exp.start_export(StartDocxExportCommand(
        project_ref=w["project"], draft_ref=draft.draft_ref,
        operator_ref="U1", idempotency_key="e-c2",
    ))
    session.commit()
    exp.confirm_baseline(ConfirmBaselineCommand(
        project_ref=w["project"], export_ref=result.export_ref,
        operator_ref="U1", idempotency_key="b-c2",
    ))
    session.commit()
    read2 = _svc(session).read_asset_fragment(w["project"], "chart", w["chart_ok"])
    assert read2.baseline_ref is not None  # 已冻结为发布基线


# ---- 追溯图：chart→document 派生承接边 ----

def test_chart_document_reference_edge_derived(session):
    from app.repositories.sqlalchemy import SqlIssueRepository, SqlTraceLinkRepository
    from app.repositories.trace_read import TraceReadRepository
    from app.services.trace_analysis import TraceAnalysisService

    w = _seed(session)
    _draft(session, w)
    svc = TraceAnalysisService(
        TraceReadRepository(session), SqlTraceLinkRepository(session), SqlIssueRepository(session),
    )
    chain = svc.read_chain(w["project"], focus_type="chart", focus_ref=w["chart_ok"],
                           direction="downstream", depth=1)
    edges = [e for level in chain.levels for e in level.edges]
    doc_edges = [e for e in edges if e.relation_kind == "document_reference"]
    assert len(doc_edges) == 1
    assert doc_edges[0].origin == "derived"
    nodes = [n for level in chain.levels for n in level.nodes]
    assert any(n.node_type == "document" for n in nodes)
