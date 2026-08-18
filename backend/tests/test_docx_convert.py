"""docx 转换器·图形围栏栅格化单测（SCN-005 图形渲染修复）。

覆盖：mermaid/plantuml 围栏 → 内嵌图片；普通代码围栏保持等宽源码；
渲染不可用/失败时降级为源码块（绝不因单图失败丢内容）。
"""
from pathlib import Path

from docx import Document as DocxDocument

from app.adapters import docx_convert
from app.adapters.diagram_render import DiagramRenderUnavailable

_BINDING = {"body_size_pt": 12, "heading_sizes_pt": {"1": 16}}
_META = {"title": "需求规格说明", "version": "V1.0"}


def _convert(md: str, tmp_path: Path):
    out = docx_convert.convert_markdown_to_docx(md, tmp_path / "d.docx", _BINDING, _META)
    return DocxDocument(str(out))


def test_plantuml_and_plain_code_coexist(tmp_path):
    md = (
        "# 1 概述\n\n正文一段。\n\n"
        "```plantuml\n@startuml\nAlice -> Bob: 下单\n@enduml\n```\n\n"
        "```python\nprint('hello')\n```\n"
    )
    doc = _convert(md, tmp_path)
    texts = [p.text for p in doc.paragraphs]
    assert doc.inline_shapes and any("PICTURE" in str(s.type) for s in doc.inline_shapes)
    assert not any("startuml" in t for t in texts)  # 图形源码不外泄
    assert any("print('hello')" in t for t in texts)  # 普通代码仍作源码保留


def test_render_failure_falls_back_to_source(tmp_path, monkeypatch):
    def _boom(source: str, fmt: str):
        raise DiagramRenderUnavailable("tool missing")

    monkeypatch.setattr(docx_convert, "render_to_png", _boom)
    md = "# 1\n\n```mermaid\nflowchart LR\n A --> B\n```\n"
    doc = _convert(md, tmp_path)
    texts = [p.text for p in doc.paragraphs]
    assert not doc.inline_shapes  # 渲染失败：不产图
    assert any("flowchart LR" in t for t in texts)  # 降级为源码，内容不丢


def test_list_form_heading_sizes_convert_without_error(tmp_path):
    """定制器登记的模板 heading_sizes_pt 是列表形态（[16,14,13]），转换器须按位次归一。

    2026-07-26 生产故障：只认字典形态时，全部模板定制器产出的模板在导出一步炸
    AttributeError，被兜底 except 吞成「转换发生未预期错误」。列表按位次映射 1..n 级。
    """
    binding = {"body_size_pt": 12, "heading_sizes_pt": [16, 14, 13]}
    md = "# 1 概述\n\n## 1.1 目标\n\n### 细则\n\n正文。\n"
    out = docx_convert.convert_markdown_to_docx(md, tmp_path / "d.docx", binding, _META)
    doc = DocxDocument(str(out))
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(headings) == 3
    # 位次映射生效：一级 16pt、二级 14pt、三级 13pt
    got = [h.runs[0].font.size.pt for h in headings]
    assert got == [16.0, 14.0, 13.0]


def test_list_form_heading_sizes_deeper_level_falls_back(tmp_path):
    """列表只给三级时，第四级标题回落默认 13pt，不越界取值。"""
    binding = {"body_size_pt": 12, "heading_sizes_pt": [16, 14]}
    md = "# 1\n\n### 深层\n\n正文。\n"
    out = docx_convert.convert_markdown_to_docx(md, tmp_path / "d2.docx", binding, _META)
    doc = DocxDocument(str(out))
    h3 = [p for p in doc.paragraphs if p.style.name == "Heading 3"]
    assert h3 and h3[0].runs[0].font.size.pt == 13.0
