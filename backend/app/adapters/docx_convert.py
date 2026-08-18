"""文档转换适配器（SCN-005-P03-N05）：Markdown 定稿版本 → 候选 docx 导出件。

只做格式化转换，不生成需求语义、不写内部仓储（失败以 ConversionError 抛出）。
排版由模板导出绑定驱动：标题黑体逐级、正文宋体/Times New Roman、
**正文段落首行缩进 2 字符**（Markdown 源无缩进，docx 必须有）。
标记 <!--convert-fail--> 为确定性失败注入（验收/演示转换失败分支用）。
"""
from __future__ import annotations

import re
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.adapters.diagram_render import (
    DiagramRenderError,
    DiagramRenderUnavailable,
    png_size,
    render_to_png,
)

FAIL_MARKER = "<!--convert-fail-->"
# 正文可用宽度约 6 英寸（A4/Letter 默认页边距）；栅格图封顶到此宽，超宽等比缩放。
_MAX_IMG_WIDTH_INCHES = 6.0


class ConversionError(Exception):
    """转换失败（原因随消息返回；不改变任何内部治理事实）。"""


def _set_run_fonts(run, east_asia: str, ascii_font: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def _indent_first_line_chars(paragraph, chars: int, body_size_pt: float) -> None:
    """首行缩进 N 字符：写 w:firstLineChars（字符单位，中文排版权威），并落 firstLine 兜底。"""
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(body_size_pt * chars)  # 兜底：不支持 firstLineChars 的渲染器
    ind = paragraph._p.get_or_add_pPr().find(qn("w:ind"))
    if ind is not None:
        ind.set(qn("w:firstLineChars"), str(chars * 100))  # 100 = 1 字符


_BOLD_SPLIT = re.compile(r"(\*\*.+?\*\*)")
_MONO_FONT = "Consolas"


def _add_code_paragraph(doc, text: str, binding: dict):
    """围栏代码行（图表源码等）：等宽、无首行缩进、紧凑行距，保留行内空白。"""
    size = float(binding.get("body_size_pt", 12)) - 1.5
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text if text else " ")
    _set_run_fonts(run, binding.get("body_font_east_asia", "宋体"), _MONO_FONT, size)
    return p


def _add_diagram_image(doc, source: str, fmt: str) -> bool:
    """图形围栏渲染为居中图片；渲染不可用/失败返回 False，由调用方降级为源码块（绝不丢内容）。"""
    try:
        png = render_to_png(source, fmt)
    except (DiagramRenderError, DiagramRenderUnavailable):
        return False
    w_px, _ = png_size(png)
    # mermaid 以 -s 2 出图（2 倍像素），按 96dpi 反算原始宽度再封顶；plantuml 为原生像素。
    css_px = (w_px / 2) if fmt == "mermaid" else w_px
    native_in = (css_px / 96.0) if css_px else _MAX_IMG_WIDTH_INCHES
    width = Inches(min(native_in, _MAX_IMG_WIDTH_INCHES))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(BytesIO(png), width=width)
    return True


def _emit_fence(doc, lang: str, lines: list[str], binding: dict) -> None:
    """闭合围栏出块：mermaid/plantuml 优先渲染为图片，其余（或渲染失败）逐行等宽源码。"""
    if lang in ("mermaid", "plantuml") and _add_diagram_image(doc, "\n".join(lines), lang):
        return
    for text in lines:
        _add_code_paragraph(doc, text, binding)


def _add_body_paragraph(doc, text: str, binding: dict, indent: bool = True):
    body_ea = binding.get("body_font_east_asia", "宋体")
    body_ascii = binding.get("body_font_ascii", "Times New Roman")
    size = float(binding.get("body_size_pt", 12))
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = float(binding.get("line_spacing", 1.5))
    for part in _BOLD_SPLIT.split(text):
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        run = p.add_run(part[2:-2] if bold else part)
        _set_run_fonts(run, body_ea, body_ascii, size, bold=bold)
    if indent:
        _indent_first_line_chars(p, int(binding.get("first_line_indent_chars", 2)), size)
    return p


def _add_heading(doc, text: str, level: int, binding: dict):
    # heading_sizes_pt 在库内存在两种形态：内置模板是层级字典 {"1":16,"2":14,"3":13}，
    # 模板定制器登记的模板是按层级排列的列表 [16,14,13]（登记校验不核该字段形态，
    # 存量两种并存）。此处统一归一为字典口径，列表按位次映射到 1..n 级——只认字典会让
    # 全部定制器模板在导出时炸 AttributeError（2026-07-26 生产故障）。
    raw_sizes = binding.get("heading_sizes_pt") or {}
    if isinstance(raw_sizes, list):
        sizes = {str(i + 1): v for i, v in enumerate(raw_sizes)}
    else:
        sizes = {str(k): v for k, v in raw_sizes.items()}
    size = float(sizes.get(str(level), 13))
    head_ea = binding.get("heading_font_east_asia", "黑体")
    head_ascii = binding.get("heading_font_ascii", "Arial")
    p = doc.add_heading(level=min(level, 9))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_fonts(run, head_ea, head_ascii, size, bold=True)
    run.font.color.rgb = RGBColor(0, 0, 0)  # 去掉内置 Heading 蓝色，正式文档用黑色
    return p


_TABLE_SEPARATOR = re.compile(r"^[-: ]+$")


def _flush_table(doc, buffered_rows: list[str], binding: dict) -> None:
    """把连续的 Markdown 表格行渲染为 docx 表格（条目属性表 / markdown_table 图表）。"""
    parsed: list[list[str]] = []
    for raw in buffered_rows:
        cells = [c.strip() for c in raw.strip().strip("|").split("|")]
        if cells and all(_TABLE_SEPARATOR.fullmatch(c) for c in cells if c):
            continue  # 分隔行不输出
        parsed.append(cells)
    if not parsed:
        return
    cols = max(len(row) for row in parsed)
    table = doc.add_table(rows=len(parsed), cols=cols)
    table.style = "Table Grid"
    body_ea = binding.get("body_font_east_asia", "宋体")
    body_ascii = binding.get("body_font_ascii", "Times New Roman")
    size = float(binding.get("body_size_pt", 12)) - 1.5
    for row_index, row in enumerate(parsed):
        for col_index in range(cols):
            text = row[col_index] if col_index < len(row) else ""
            paragraph = table.cell(row_index, col_index).paragraphs[0]
            for part in _BOLD_SPLIT.split(text):
                if not part:
                    continue
                bold = part.startswith("**") and part.endswith("**")
                run = paragraph.add_run(part[2:-2] if bold else part)
                _set_run_fonts(run, body_ea, body_ascii, size, bold=bold or row_index == 0)


def _add_cover(doc, meta: dict, binding: dict) -> None:
    title = meta.get("title", "需求规格说明")
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_run_fonts(run, binding.get("heading_font_east_asia", "黑体"),
                   binding.get("heading_font_ascii", "Arial"), 26, bold=True)
    for text in (meta.get("project_name", ""), f"版本：{meta.get('version', 'V1.0')}",
                 f"日期：{meta.get('date', date.today().isoformat())}"):
        if not text:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_run_fonts(run, binding.get("body_font_east_asia", "宋体"),
                       binding.get("body_font_ascii", "Times New Roman"), 14)
    doc.add_page_break()


def convert_markdown_to_docx(markdown: str, out_path: Path, binding: dict, meta: dict) -> Path:
    """把 Markdown 定稿内容渲染为 docx。标题→Heading 样式；正文段→首行缩进 2 字符。"""
    if FAIL_MARKER in markdown:
        raise ConversionError("文档转换失败：内容包含失败注入标记（演示/验收用）")
    if not markdown.strip():
        raise ConversionError("文档转换失败：Markdown 定稿内容为空")

    doc = Document()
    _add_cover(doc, meta, binding)

    in_fence = False
    fence_lang = ""
    fence_buf: list[str] = []
    table_rows: list[str] = []
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if not in_fence:  # 开栏：记语言标签（图形围栏据此渲染），缓冲栏内内容
                if table_rows:
                    _flush_table(doc, table_rows, binding)
                    table_rows = []
                in_fence = True
                fence_lang = stripped[3:].strip().lower()
                fence_buf = []
            else:  # 闭栏：整块出图/出源码（围栏行本身不进正文）
                _emit_fence(doc, fence_lang, fence_buf, binding)
                in_fence = False
                fence_lang = ""
                fence_buf = []
            continue
        if in_fence:
            fence_buf.append(line)
            continue
        if stripped.startswith("|"):
            table_rows.append(stripped)
            continue
        if table_rows:
            _flush_table(doc, table_rows, binding)
            table_rows = []
        if not stripped or stripped.startswith("<!--"):
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            _add_heading(doc, m.group(2).strip(), len(m.group(1)), binding)
            continue
        if re.match(r"^[-*]\s+", stripped):
            p = _add_body_paragraph(doc, re.sub(r"^[-*]\s+", "", stripped), binding, indent=False)
            p.style = doc.styles["List Bullet"]
            continue
        if stripped.startswith(">"):
            _add_body_paragraph(doc, stripped.lstrip("> "), binding, indent=False)
            continue
        _add_body_paragraph(doc, stripped, binding, indent=True)
    if in_fence:  # 未闭合围栏兜底出块
        _emit_fence(doc, fence_lang, fence_buf, binding)
    if table_rows:
        _flush_table(doc, table_rows, binding)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
